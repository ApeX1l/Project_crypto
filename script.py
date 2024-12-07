import csv
import os
import sys
from pathlib import Path
import sqlite3
import pandas as pd
import pandas_ta as ta
from PyQt6.QtSql import QSqlDatabase, QSqlTableModel
from PyQt6 import uic
import pyqtgraph.exporters
from PyQt6.QtCore import Qt, QPoint, pyqtSignal
from PyQt6.QtGui import QColor, QPen, QPainter, QPixmap, QFont, QIcon
import pyqtgraph as pq
from PyQt6.QtWidgets import QApplication, QFileDialog, QLabel, QWidget, QTextEdit, QPushButton, QColorDialog, \
    QDoubleSpinBox, QCheckBox, QComboBox
from PyQt6.QtWidgets import QMainWindow
from docx import Document
from docx.shared import Mm
from PIL import Image
from fpdf import FPDF


class Main(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('new.ui', self)
        self.setWindowTitle('Анализ криптовалюты')
        self.setWindowIcon(QIcon('main_pic2.png'))
        self.setFixedSize(1146, 763)
        with open('notepad_settings.txt', 'w') as f:
            f.write('9.0, False, Arial')
        with open('main_settings.txt', 'w', encoding='utf-8') as f:
            f.write('Желтый, Белый, Зеленый, Красный, Черный, Светлый, Серый')

        self.setStyleSheet("QMainWindow { background-color: rgb(250, 248, 245);}")
        self.last_point = QPoint()

        self.graphicsView.setBackground('black')
        self.font_size.setRange(0.1, 2)
        self.font_size.setDecimals(1)
        self.font_size.setSingleStep(0.1)

        self.drawing = False
        self.start_point = QPoint()
        self.end_point = QPoint()

        self.clozze = None

        self.value_tech.setRange(1, 20)
        self.build_tech = None

        self.pen = QColor(255, 255, 255)  # Черный цвет, толщина 3 пикселя

        self.initUI()

    def initUI(self):
        self.coords = QLabel(self)  # координаты зажатой левой кнопки мышки
        self.coords.setText("Координаты: None, None")
        self.coords.move(30, 600)
        self.coords.setStyleSheet('background-color: rgb(250, 165, 37)')
        self.coords.resize(self.coords.sizeHint())

        self.loadButton.clicked.connect(self.load)  # кнопка загрузки графика

        self.brushButton.setCheckable(True)  # кнопка для режима рисования
        self.brushButton.clicked[bool].connect(self.paint_brush)

        self.notepad.clicked.connect(self.notepad_window)  # кнопка для открытия блокнота

        self.colorButton.setStyleSheet('background: rgb(255, 255, 255)')  # кнопка для установки цвета рисования
        self.colorButton.clicked.connect(self.color_brush)

        self.choiceButton.clicked.connect(self.tech)  # кнопка для построения технических индикаторов

        self.parametresButton.clicked.connect(self.main_parametres)

        self.export_docx.clicked.connect(self.save_file)
        self.export_pdf.clicked.connect(self.save_file)
        self.export_png.clicked.connect(self.save_file)

        self.min.stateChanged.connect(self.build)
        self.max.stateChanged.connect(self.build)
        self.open.stateChanged.connect(self.build)
        self.cloze.stateChanged.connect(self.build)

        self.resetButton.clicked.connect(self.reset)

    def load(self):  # загрузка графика
        download_path = os.path.join(os.path.expanduser("~"), "Downloads")
        fname = QFileDialog.getOpenFileName(self, 'Выбрать файл с данными', download_path,
                                            'Данные(*.csv);; Данные(*.txt)')[0]
        if fname:
            self.new_load()
            with open(fname, mode='r', encoding='utf-8') as f:
                self.reader = list(csv.DictReader(f, delimiter=';'))
                date = [str(i['<DATE>']) for i in self.reader]
                maxx = [float(i['<HIGH>']) for i in self.reader]
                minn = [float(i['<LOW>']) for i in self.reader]
                openn = [float(i['<OPEN>']) for i in self.reader]
                cloze = [float(i['<CLOSE>']) for i in self.reader]
                volume = [float(i['<VOL>']) for i in self.reader]
                count_bars = [int(i) for i in range(len(self.reader))]
                self.clozze = self.graphicsView.plot(count_bars, cloze)
                self.minn = self.graphicsView.plot(count_bars, minn, pen='g')
                self.maxx = self.graphicsView.plot(count_bars, maxx, pen='r')
                self.openum = self.graphicsView.plot(count_bars, openn, pen='y')
                self.minn.hide()
                self.maxx.hide()
                self.openum.hide()
                self.graphicsView.getViewBox().setLimits(xMin=0, xMax=len(self.reader))
                self.cloze.setChecked(True)
                try:
                    connection = sqlite3.connect('my_database.db')
                    cursor = connection.cursor()
                    cursor.execute('DROP TABLE Prices')
                    cursor.execute('''
                            CREATE TABLE IF NOT EXISTS Prices (
                            Id INTEGER PRIMARY KEY,
                            Date TEXT,
                            Open REAL,
                            High REAL,
                            Low REAL,
                            Close REAL,
                            Volume REAL
                            )
                            ''')
                    data = list(zip(date, openn, maxx, minn, cloze, volume))
                    cursor.executemany(
                        'INSERT INTO Prices (date, open, high, low, close, volume) VALUES (?, ?, ?, ?, ?, ?)',
                        data)
                    connection.commit()
                    connection.close()
                    db = QSqlDatabase.addDatabase('QSQLITE')
                    db.setDatabaseName('my_database.db')
                    db.open()
                    model = QSqlTableModel(self, db)
                    model.setTable('Prices')
                    model.select()
                    self.tableView.setModel(model)
                except Exception:
                    pass

    def new_load(self):
        if self.clozze is not None:
            self.clozze.clear()
            self.minn.clear()
            self.maxx.clear()
            self.openum.clear()
            if self.build_tech is not None:
                self.build_tech.clear()

    def build(self):
        tr = self.sender().text()
        try:
            if self.sender().isChecked():
                if tr == 'Цена минимума':
                    self.minn.show()
                elif tr == 'Цена максимума':
                    self.maxx.show()
                elif tr == 'Цена открытия':
                    self.openum.show()
                elif tr == 'Цена закрытия':
                    self.clozze.show()
            else:
                if tr == 'Цена минимума':
                    self.minn.hide()
                elif tr == 'Цена максимума':
                    self.maxx.hide()
                elif tr == 'Цена открытия':
                    self.openum.hide()
                elif tr == 'Цена закрытия':
                    self.clozze.hide()
        except Exception:
            pass

    def tech(self):
        try:
            if self.build_tech is not None:
                self.build_tech.clear()
            indicator = self.technical_indikactor.currentText()
            all_price_close = [float(i['<CLOSE>']) for i in self.reader]  # цены закрытия
            count_bars = [int(i) for i in range(len(self.reader))]  # кол-во баров
            size = self.value_tech.value()  # значение окна
            data = {'Close': all_price_close}  # получение данных для построения
            df = pd.DataFrame(data)  # построение таблицы
            if indicator == 'Скользящая средняя(SMA)':
                df.ta.sma(close='Close', length=size, append=True)
                sma_values = df[f'SMA_{size}'].dropna().tolist()
                self.build_tech = self.graphicsView.plot(count_bars[size - 1:], sma_values, pen='gray')
            elif indicator == 'Относительный индекс силы(RSI)':
                df.ta.rsi(close='Close', length=size, append=True)
                sma_values = df[f'RSI_{size}'].dropna().tolist()
                self.build_tech = self.graphicsView.plot(sma_values, pen='gray')
            elif indicator == 'Скользящая средняя затрат(EMA)':
                df.ta.ema(close='Close', length=size, append=True)
                sma_values = df[f'EMA_{size}'].dropna().tolist()
                self.build_tech = self.graphicsView.plot(count_bars[size - 1:], sma_values, pen='gray')
            elif indicator == 'Балансовый объем(OBV)':
                volume = [float(i['<VOL>']) for i in self.reader]
                data = {'Close': all_price_close, 'Volume': volume}
                df = pd.DataFrame(data)
                obv_values = [0.0]  # Инициализация
                for i in range(1, len(df)):
                    if df['Close'][i] > df['Close'][i - 1]:
                        obv_values.append(df['Volume'][i])
                    elif df['Close'][i] < df['Close'][i - 1]:
                        obv_values.append(-df['Volume'][i])
                    else:
                        obv_values.append(0.0)
                obv_values = pd.Series(obv_values).cumsum()
                self.build_tech = self.graphicsView.plot(obv_values, pen='gray')
            self.build_tech.show()
        except Exception:
            pass

    def reset(self):
        try:
            self.build_tech.hide()
            self.build_tech.clear()
        except Exception:
            pass

    def paint_brush(self, pressed):  # переход в режим рисования
        if pressed:
            self.show_pixmap = True
            exporter = pq.exporters.ImageExporter(self.graphicsView.plotItem)
            exporter.params.param('width').setValue(701, blockSignal=exporter.widthChanged)
            exporter.params.param('height').setValue(321, blockSignal=exporter.heightChanged)
            exporter.export('fileName.png')
            self.graphicsView.hide()
            self.pixmap = QPixmap('fileName.png')
            self.repaint()
        else:
            self.graphicsView.show()
            self.show_pixmap = False
            self.update()

    def color_brush(self):  # выбор цвета кисти
        color = QColorDialog.getColor()
        if color.isValid():
            self.colorButton.setStyleSheet(f'background-color: {color.name()}')
            r, g, b, a = color.getRgb()
            self.pen = QColor(r, g, b)

    def paintEvent(self, event):
        font = self.font_size.value()
        try:
            painter = QPainter(self)
            if self.show_pixmap:  # условие при котором можно рисовать
                painter.drawPixmap(190, 130, self.pixmap)
                pen = QPen(self.pen, font)
                self.check = self.choiceGroup.checkedButton().text()
                if self.drawing:  # рисование линий
                    if self.check == 'Линейка':
                        painter.setPen(pen)
                        painter.drawLine(self.start_point, self.end_point)
        except Exception:
            pass

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            try:
                if self.check == 'Линейка':
                    self.start_point = event.position().toPoint()  # ровная линия
                    self.end_point = self.start_point
                else:
                    self.last_point = event.position().toPoint()  # свободное рисование
            except Exception:
                pass
            self.drawing = True

    def mouseMoveEvent(self, event):
        try:
            self.coords.setText(f"Координаты: {event.pos().x()}, {event.pos().y()}")
            if self.drawing:
                if self.check == 'Кисть':
                    font = self.font_size.value()
                    painter = QPainter(self.pixmap)
                    pen = QPen(self.pen, font)
                    painter.setPen(pen)
                    self.last_point = QPoint(self.last_point.x() - 190, self.last_point.y() - 130)
                    painter.drawLine(self.last_point, QPoint(event.pos().x() - 190, event.pos().y() - 130))
                    self.last_point = event.position().toPoint()
                else:
                    self.end_point = event.position().toPoint()
                self.update()
        except Exception:
            pass

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton and self.drawing:
            try:
                if self.check == 'Линейка':
                    self.end_point = event.position().toPoint()
                    self.draw_line()
            except Exception:
                pass
            self.drawing = False

    def draw_line(self):
        try:
            font = self.font_size.value()
            painter = QPainter(self.pixmap)
            pen = QPen(self.pen, font)
            painter.setPen(pen)
            self.start_point = QPoint(self.start_point.x() - 190, self.start_point.y() - 130)
            self.end_point = QPoint(self.end_point.x() - 190, self.end_point.y() - 130)
            painter.drawLine(self.start_point, self.end_point)
            self.update()
        except Exception:
            pass

    def notepad_window(self):
        self.note = Notepad()
        self.note.show()

    def main_parametres(self):
        self.parametres = Main_parameters()
        self.parametres.applied.connect(self.apply_settings)
        self.parametres.show()

    def apply_settings(self, open_box, close_box, min_box, max_box, graph_box, back, tech_box):
        try:
            tr = self.colorgraph(open_box)
            self.openum.setPen(tr)
            tr = self.colorgraph(close_box)
            self.clozze.setPen(tr)
            tr = self.colorgraph(min_box)
            self.minn.setPen(tr)
            tr = self.colorgraph(max_box)
            self.maxx.setPen(tr)
            tr = self.colorgraph(graph_box)
            self.graphicsView.setBackground(tr)
            tr = self.colorgraph(tech_box)
            self.build_tech.setPen(tr)
        except Exception:
            pass
        if back == 'Темный':
            self.setStyleSheet("QMainWindow { background-color: rgb(30, 28, 26);}")
        else:
            self.setStyleSheet("QMainWindow { background-color: rgb(250, 248, 245);}")

    def colorgraph(self, name):
        if name == 'Белый':
            return 'w'
        elif name == 'Черный':
            return 'black'
        elif name == 'Желтый':
            return 'y'
        elif name == 'Синий':
            return 'blue'
        elif name == 'Красный':
            return 'red'
        elif name == 'Зеленый':
            return 'green'
        elif name == 'Серый':
            return 'gray'

    def save_file(self):
        try:
            date = [str(i['<DATE>']) for i in self.reader]
            maxx = [float(i['<HIGH>']) for i in self.reader]
            minn = [float(i['<LOW>']) for i in self.reader]
            openn = [float(i['<OPEN>']) for i in self.reader]
            cloze = [float(i['<CLOSE>']) for i in self.reader]
            volume = [float(i['<VOL>']) for i in self.reader]
            idd = [int(i) + 1 for i in range(len(self.reader))]
            data = [idd, date, maxx, minn, openn, cloze, volume]
            name_cols = ['Id', 'Date', 'Open', 'High', 'Low', 'Close', 'Volume']
            count_rows = len(self.reader) + 1
            self.make_screenshot()
            if self.sender().text() == 'docx':
                fname = QFileDialog.getSaveFileName(self, 'Выбрать файл куда сохранить', '', 'Формат word(*.docx)')[0]
                if fname:
                    document = Document()
                    document.add_picture('save.png', width=Mm(180))
                    table = document.add_table(rows=count_rows, cols=7, style='Table Grid')
                    for i in range(7):
                        table.cell(0, i).text = name_cols[i]
                    for i, row_data in enumerate(zip(*data)):
                        row = table.rows[i + 1]
                        for j, cell_data in enumerate(row_data):
                            row.cells[j].text = str(cell_data)
                    document.save(fname)
            elif self.sender().text() == 'pdf':
                fname = QFileDialog.getSaveFileName(self, 'Выбрать файл куда сохранить', '', 'Формат pdf(*.pdf)')[0]
                if fname:
                    data = list(zip(idd, date, maxx, minn, openn, cloze, volume))
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font('helvetica', size=10)
                    pdf.image('save.png', w=190)
                    for i in range(7):
                        pdf.cell(w=25, text=name_cols[i])
                    pdf.ln()
                    for row in data:
                        for i, cell in enumerate(row):
                            pdf.cell(w=25, h=10, text=str(cell), border=1, align='C')
                        pdf.ln()
                    pdf.output(fname)
            else:
                fname = QFileDialog.getSaveFileName(self, 'Выбрать файл куда сохранить', '', 'Формат png(*.png)')[0]
                if fname:
                    img = Image.open('save.png')
                    img.save(fname)
        except Exception as e:
            pass

    def make_screenshot(self):
        x = 190
        y = 130
        width = 701
        height = 321
        try:
            pixmap = QPixmap(self.size())
            painter = QPainter(pixmap)
            self.render(painter)
            painter.end()
            image = Image.fromqpixmap(pixmap)
            cropped_image = image.crop((x, y, x + width, y + height))
            cropped_image.save('save.png')
        except Exception:
            pass


class Main_parameters(QWidget):
    applied = pyqtSignal(str, str, str, str, str, str, str)

    def __init__(self):
        super().__init__()
        uic.loadUi('main_settings.ui', self)
        self.setWindowTitle('Настройки главного окна')
        self.setFixedSize(592, 313)
        self.setWindowIcon(QIcon('settings.png'))
        self.resultButton.clicked.connect(self.result)
        self.load_settings()

    def result(self):
        self.applied.emit(self.open_box.currentText(), self.close_box.currentText(), self.min_box.currentText(),
                          self.max_box.currentText(), self.graph_box.currentText(),
                          self.background_box.currentText(), self.tech_box.currentText())
        self.save_settings()
        self.close()

    def save_settings(self):
        with open('main_settings.txt', 'w', encoding='utf-8') as f:
            f.write(f'{self.open_box.currentText()}, {self.close_box.currentText()}, {self.min_box.currentText()}, '
                    f'{self.max_box.currentText()}, {self.graph_box.currentText()}, '
                    f'{self.background_box.currentText()}, {self.tech_box.currentText()}')

    def load_settings(self):
        path = Path('main_settings.txt')
        if path.exists():
            with open('main_settings.txt', 'r', encoding='utf-8') as f:
                try:
                    open_box, close_box, min_box, max_box, graph_box, back, tech = f.readline().strip().split(', ')
                    self.open_box.setCurrentText(open_box)
                    self.close_box.setCurrentText(close_box)
                    self.min_box.setCurrentText(min_box)
                    self.max_box.setCurrentText(max_box)
                    self.graph_box.setCurrentText(graph_box)
                    self.background_box.setCurrentText(back)
                    self.tech_box.setCurrentText(tech)
                except Exception:
                    pass


class Notepad(QWidget):  # блокнот
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Блокнот')
        self.setWindowIcon(QIcon('notepad.png'))
        self.setFixedSize(450, 450)
        self.textEdit = QTextEdit(self)
        self.textEdit.move(0, 50)
        self.textEdit.resize(450, 400)

        self.button_save = QPushButton('Сохранить файл', self)
        self.button_save.resize(self.button_save.sizeHint())
        self.button_save.move(300, 10)
        self.button_save.clicked.connect(self.save)

        self.button_open = QPushButton('Открыть файл', self)
        self.button_open.resize(self.button_open.sizeHint())
        self.button_open.move(75, 10)
        self.button_open.clicked.connect(self.open)

        self.button_parameters = QPushButton('Настройки', self)
        self.button_parameters.resize(self.button_parameters.sizeHint())
        self.button_parameters.move(195, 10)
        self.button_parameters.clicked.connect(self.parametres)

    def parametres(self):
        self.notepad_parameters = Notepad_parameters()
        self.notepad_parameters.applied.connect(self.apply_settings)
        self.notepad_parameters.show()

    def apply_settings(self, font_size, font_bold, font_name):
        self.textEdit.setCurrentFont(QFont(font_name))
        self.textEdit.setFontPointSize(font_size)
        if font_bold:
            self.textEdit.setFontWeight(QFont.Weight.Bold)
        else:
            self.textEdit.setFontWeight(QFont.Weight.Normal)

    def save(self):  # сохранение записей
        fname = QFileDialog.getSaveFileName(self, 'Выбрать файл куда сохранить', '', 'Текстовые файлы (*.txt)')[0]
        if fname:
            with open(fname, 'w') as file:
                data = self.textEdit.toPlainText()
                file.write(data)

    def open(self):  # открыть файл
        fname = QFileDialog.getOpenFileName(self, 'Выбрать файл с данными', '', 'Текстовые файлы (*.txt)')[0]
        if fname:
            with open(fname, 'r') as file:
                data = file.read()
                self.textEdit.setText(data)


class Notepad_parameters(QWidget):
    applied = pyqtSignal(float, bool, str)

    def __init__(self):
        super().__init__()
        self.setFixedSize(300, 300)
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Параметры блокнота')
        self.setStyleSheet("QWidget { background-color: white;}")
        self.setWindowIcon(QIcon('settings.png'))

        self.resultButton = QPushButton('Применить', self)
        self.resultButton.resize(self.resultButton.sizeHint())
        self.resultButton.move(115, 250)
        self.resultButton.clicked.connect(self.result)

        self.label = QLabel('Размер шрифта', self)
        self.label.resize(150, 50)
        self.label.move(10, 10)
        self.label.setStyleSheet("""
          QLabel {
            font-weight: bold;
          }
        """)

        self.font_size = QDoubleSpinBox(self)
        self.font_size.setRange(2, 20)
        self.font_size.setDecimals(1)
        self.font_size.setSingleStep(1)
        self.font_size.move(15, 50)

        self.font_bold = QCheckBox('Жирный шрифт', self)
        self.font_bold.setStyleSheet("""QCheckBox { font-weight: bold; }""")
        self.font_bold.move(140, 50)

        self.font_name = QComboBox(self)
        self.font_name.addItems(['Arial', 'Agency FB', 'Bahnschrift SemiCondensed',
                                 'MS Shell Dlg 2', 'Lucida Calligraphy', 'Rockwell Extra Bold', 'Myanmar Text'])
        self.font_name.move(80, 120)

        self.load_settings()

    def result(self):
        self.applied.emit(self.font_size.value(), self.font_bold.isChecked(), self.font_name.currentText())
        self.save_settings()
        self.close()

    def load_settings(self):
        path = Path('notepad_settings.txt')
        if path.exists():
            with open('notepad_settings.txt', 'r') as f:
                try:
                    font_size, font_bold, font_name = f.readline().strip().split(', ')
                    self.font_size.setValue(float(font_size))
                    self.font_bold.setChecked(True if font_bold == 'True' else False)
                    self.font_name.setCurrentText(font_name)
                except Exception as e:
                    print(e)

    def save_settings(self):
        with open('notepad_settings.txt', 'w') as f:
            f.write(f'{self.font_size.value()}, {self.font_bold.isChecked()}, {self.font_name.currentText()}')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Main()
    ex.show()
    sys.exit(app.exec())
